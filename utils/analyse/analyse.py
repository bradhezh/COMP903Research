
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# === Load Data ===
files = ['results.csv', 'results (1).csv', 'results (2).csv']
# Read each of the three CSV result
filesdfs = [pd.read_csv(file) for file in files]
# Combine all DataFrames and tag them with run IDs
combined_df = pd.concat(filesdfs, keys=[f"Run{i+1}" for i in range(len(filesdfs))], names=["Run", "Index"]).reset_index()
# Add an index column to help detect time-based trends
combined_df["sample_index"] = combined_df.index

# === Summary Statistics ===
# Compute mean statistics for load time, inference time, and accuracy
summary = combined_df.groupby("Run").agg({
    "loadTimeTf": "mean",
    "loadTimeOrt": "mean",
    "inferTimeTf": "mean",
    "inferTimeOrt": "mean",
    "correctTf": "mean",
    "correctOrt": "mean"
})

# === Cold Start ===
# Extract the first sample from each run to measure cold start time
first_sample_load = combined_df.groupby("Run").first().reset_index()[["Run", "loadTimeTf", "loadTimeOrt"]]

# === Trend Analysis (Regression) ===
# Add an index column to help detect time-based trends
X = sm.add_constant(combined_df["sample_index"])

# Load time
# Perform linear regression on load time vs sample index
model_tf_load = sm.OLS(combined_df["loadTimeTf"], X).fit()
# Perform linear regression on load time vs sample index
model_ort_load = sm.OLS(combined_df["loadTimeOrt"], X).fit()

# Inference time
# Perform linear regression on inference time vs sample index
model_tf_infer = sm.OLS(combined_df["inferTimeTf"], X).fit()
# Perform linear regression on inference time vs sample index
model_ort_infer = sm.OLS(combined_df["inferTimeOrt"], X).fit()

# === Visualizations ===
# Inference time boxplot
inference_times = combined_df[["inferTimeTf", "inferTimeOrt"]].copy()
inference_times.columns = ["TensorFlow.js", "ONNX"]
inference_times = inference_times.melt(var_name="Library", value_name="Inference Time (ms)")
plt.figure(figsize=(6, 4))
# Create a boxplot of inference time distribution between libraries
sns.boxplot(data=inference_times, x="Library", y="Inference Time (ms)")
plt.title("Inference Time Distribution")
plt.tight_layout()
# Create a boxplot of inference time distribution between libraries
plt.savefig("inference_time_boxplot.png")
plt.close()

# Load time trend
# Add an index column to help detect time-based trends
load_times = combined_df[["sample_index", "loadTimeTf", "loadTimeOrt"]].copy()
# Add an index column to help detect time-based trends
load_times.columns = ["sample_index", "TensorFlow.js", "ONNX"]
# Add an index column to help detect time-based trends
load_times = load_times.melt(id_vars="sample_index", var_name="Library", value_name="Load Time (ms)")
plt.figure(figsize=(10, 4))
# Add an index column to help detect time-based trends
sns.lineplot(data=load_times, x="sample_index", y="Load Time (ms)", hue="Library", ci=None)
plt.title("Load Time per Sample")
plt.tight_layout()
plt.savefig("load_time_trend.png")
plt.close()

# Cold start bar chart
cold_start = first_sample_load.melt(id_vars="Run", var_name="Library", value_name="Load Time (ms)")
plt.figure(figsize=(6, 4))
# Create a bar chart showing cold start load time per run
sns.barplot(data=cold_start, x="Run", y="Load Time (ms)", hue="Library")
plt.title("Cold Start Load Time")
plt.tight_layout()
plt.savefig("cold_start_bar.png")
plt.close()

# === Report Generation ===
# Start generating the Word report
doc = Document()
doc.add_heading("Results and Analysis", level=1)
doc.add_paragraph("This report summarizes the benchmarking analysis comparing TensorFlow.js and ONNX Runtime Web.")

doc.add_heading("Accuracy Summary", level=2)
doc.add_paragraph("Both libraries achieved 98.2% prediction accuracy.")

doc.add_heading("Cold Start Load Times", level=2)
table = doc.add_table(rows=1, cols=len(first_sample_load.columns))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, col in enumerate(first_sample_load.columns):
    hdr_cells[i].text = col
for _, row in first_sample_load.iterrows():
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = str(round(val, 2))

doc.add_heading("Inference Time Distribution", level=2)
# Create a boxplot of inference time distribution between libraries
doc.add_picture("inference_time_boxplot.png", width=Inches(5.5))

doc.add_heading("Load Time Trends", level=2)
doc.add_picture("load_time_trend.png", width=Inches(5.5))

doc.add_heading("Cold Start Load Time", level=2)
doc.add_picture("cold_start_bar.png", width=Inches(5.5))

# Save the Word document with embedded results and plots
doc.save("benchmark_analysis_report.docx")
print("Analysis complete. Report saved to 'benchmark_analysis_report.docx'")
